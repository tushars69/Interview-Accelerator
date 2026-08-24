"""Extract plain text from uploaded JD/resume files so the rest of the
pipeline only ever has to deal with plain strings."""
from __future__ import annotations

import io


def extract_text(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    if lower.endswith(".txt") or lower.endswith(".md"):
        return file_bytes.decode("utf-8", errors="ignore")
    # Fall back to best-effort decode for unknown types
    return file_bytes.decode("utf-8", errors="ignore")


def _extract_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def _extract_docx(file_bytes: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()
